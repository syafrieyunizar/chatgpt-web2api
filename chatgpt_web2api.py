#!/usr/bin/env python3
"""
chatgpt-web2api - ChatGPT Web to OpenAI API proxy.

Converts ChatGPT's web interface (chatgpt.com) into an OpenAI-compatible API
server.  Direct reverse-engineered approach: POST to /backend-api/conversation,
solve proof-of-work challenges, parse SSE streaming.

Usage:
    pip install curl_cffi pybase64
    python chatgpt_web2api.py [--port 6970] [--config config.json]

Client configuration (Cherry Studio, ChatBox, etc.):
    Base URL: http://localhost:6970/v1
    API Key:  your ChatGPT access token (or refresh token), or x-api-key

How it works:
    1. GET chatgpt.com page → obtain cookies (oai-did, __cf_bm, etc.)
    2. POST /backend-api/sentinel/chat-requirements (solve PoW if needed)
    3. POST /backend-api/conversation with SSE streaming
    4. Parse SSE chunks → OpenAI-compatible response (streaming or not)
"""

import json
import hashlib
import random
import re
import time
import uuid
import string
import os
import sys
import io
import argparse
from http.server import HTTPServer, BaseHTTPRequestHandler
from socketserver import ThreadingMixIn
from datetime import datetime, timedelta, timezone

try:
    from curl_cffi import requests as cffi_requests
    HAS_CFFI = True
except ImportError:
    HAS_CFFI = False

try:
    import pybase64
    HAS_PYBASE64 = True
except ImportError:
    HAS_PYBASE64 = False
    import base64 as pybase64

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    Image = None

__version__ = "1.0.0"

# ─── Configuration ───────────────────────────────────────────────────────────

DEFAULT_CONFIG = {
    "port": 6970,
    "host": "0.0.0.0",
    "host_url": "https://chatgpt.com",
    "retry_attempts": 3,
    "retry_delay_sec": 2,
    "request_timeout_sec": 120,
    "default_model": "gpt-4o-mini",
    "log_requests": True,
    "api_keys": [],
    "access_token": None,
    "refresh_token": None,
    "proxy": None,
    "history_disabled": True,
    "pow_difficulty": "0fffff",
    "impersonate": "safari15_3",
}

CONFIG = dict(DEFAULT_CONFIG)

# ─── Models ──────────────────────────────────────────────────────────────────

MODELS = {
    "gpt-5.6-luna":       {"slug": "gpt-5-6",      "desc": "GPT-5.6 Luna (latest, default)"},
    "gpt-5.5":            {"slug": "gpt-5-5",      "desc": "GPT-5.5"},
    "gpt-5.6-luna-mini":  {"slug": "gpt-5-6-mini", "desc": "GPT-5.6 Luna Mini"},
    "gpt-5.5-mini":       {"slug": "gpt-5-5-mini", "desc": "GPT-5.5 Mini"},
    "gpt-5.3-mini":       {"slug": "gpt-5-3-mini", "desc": "GPT-5.3 Mini"},
    "gpt-5.4-t-mini":     {"slug": "gpt-5-4-t-mini","desc": "GPT-5.4 Thinking Mini"},
    "gpt-4o":             {"slug": "gpt-4o",       "desc": "GPT-4o"},
    "gpt-4o-mini":        {"slug": "gpt-4o-mini",  "desc": "GPT-4o Mini"},
    "gpt-4":              {"slug": "gpt-4",        "desc": "GPT-4 (legacy)"},
    "gpt-3.5-turbo":     {"slug": "text-davinci-002-render-sha", "desc": "GPT-3.5 Turbo"},
    "o1":                 {"slug": "o1",           "desc": "o1 reasoning"},
    "o1-mini":            {"slug": "o1-mini",      "desc": "o1-mini"},
    "o1-preview":         {"slug": "o1-preview",   "desc": "o1-preview"},
    "o3":                 {"slug": "o3",           "desc": "o3 reasoning"},
    "o3-mini":            {"slug": "o3-mini",      "desc": "o3-mini"},
    "o3-mini-high":       {"slug": "o3-mini-high", "desc": "o3-mini high"},
    "research":           {"slug": "research",     "desc": "Deep Research"},
    "auto":               {"slug": "auto",         "desc": "Auto model selection"},
}

# ─── PoW (Proof of Work) ────────────────────────────────────────────────────

POW_CORES = [8, 16, 24, 32]
POW_NAVIGATOR_KEYS = [
    "webdriver−false",
    "vendor−Google Inc.",
    "cookieEnabled−true",
    "product−Gecko",
    "hardwareConcurrency−32",
    "pdfViewerEnabled−true",
]
POW_DOCUMENT_KEYS = ["location"]
POW_WINDOW_KEYS = [
    "0", "window", "self", "document", "name", "location",
    "navigator", "crypto", "localStorage", "performance",
]

_cached_scripts = []
_cached_dpl = ""
_cached_dpl_time = 0

UA = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.0 Safari/605.1.15"


def log(msg: str):
    if CONFIG["log_requests"]:
        sys.stderr.write(f"[{time.strftime('%H:%M:%S')}] {msg}\n")
        sys.stderr.flush()


def _get_parse_time():
    now = datetime.now(timezone(timedelta(hours=-5)))
    return now.strftime("%a %b %d %Y %H:%M:%S") + " GMT-0500 (Eastern Standard Time)"


def _get_pow_config(user_agent: str = UA):
    _maybe_fetch_dpl()
    return [
        random.choice([1920 + 1080, 2560 + 1440, 1920 + 1200]),
        _get_parse_time(),
        4294705152,
        0,
        user_agent,
        random.choice(_cached_scripts) if _cached_scripts else "",
        _cached_dpl,
        "en-US",
        "en-US,es-US,en,es",
        0,
        random.choice(POW_NAVIGATOR_KEYS),
        random.choice(POW_DOCUMENT_KEYS),
        random.choice(POW_WINDOW_KEYS),
        time.perf_counter() * 1000,
        str(uuid.uuid4()),
        "",
        random.choice(POW_CORES),
        time.time() * 1000 - (time.perf_counter() * 1000),
    ]


def _maybe_fetch_dpl():
    global _cached_scripts, _cached_dpl, _cached_dpl_time
    if int(time.time()) - _cached_dpl_time < 15 * 60 and _cached_dpl:
        return
    try:
        session = cffi_requests.Session(impersonate=CONFIG.get("impersonate", "safari15_3"))
        r = session.get(CONFIG["host_url"] + "/", headers={"user-agent": UA}, timeout=10)
        if r.status_code == 200:
            scripts = re.findall(r'<script[^>]+src="([^"]+)"', r.text)
            if scripts:
                _cached_scripts = scripts
            for src in scripts:
                m = re.search(r'/_next/static/([^/]+)/', src)
                if m:
                    _cached_dpl = m.group(1)
                    break
            if not _cached_dpl:
                m = re.search(r'data-build="([^"]+)"', r.text)
                if m:
                    _cached_dpl = m.group(1)
            if not _cached_dpl:
                _cached_dpl = "prod-f501fe933b3edf57aea882da888e1a544df99840"
            _cached_dpl_time = int(time.time())
            log(f"DPL fetched: {_cached_dpl}")
        session.close()
    except Exception as e:
        log(f"DPL fetch failed: {e}")
        if not _cached_dpl:
            _cached_dpl = "prod-f501fe933b3edf57aea882da888e1a544df99840"
        _cached_dpl_time = int(time.time())


def _generate_pow_answer(seed: str, diff: str, config: list) -> tuple:
    diff_len = len(diff)
    seed_enc = seed.encode()
    p1 = (json.dumps(config[:3], separators=(",", ":"), ensure_ascii=False)[:-1] + ",").encode()
    p2 = ("," + json.dumps(config[4:9], separators=(",", ":"), ensure_ascii=False)[1:-1] + ",").encode()
    p3 = ("," + json.dumps(config[10:], separators=(",", ":"), ensure_ascii=False)[1:]).encode()
    target = bytes.fromhex(diff)
    for i in range(500000):
        di = str(i).encode()
        dj = str(i >> 1).encode()
        final = p1 + di + p2 + dj + p3
        base = pybase64.b64encode(final)
        h = hashlib.sha3_512(seed_enc + base).digest()
        if h[:diff_len] <= target:
            return base.decode(), True
    return "wQ8Lk5FbGpA2NcR9dShT6gYjU7VxZ4D" + pybase64.b64encode(f'"{seed}"'.encode()).decode(), False


def _get_answer_token(seed, diff, config):
    ans, solved = _generate_pow_answer(seed, diff, config)
    return "gAAAAAB" + ans, solved


def _get_requirements_token(config):
    req, _ = _generate_pow_answer(format(random.random()), "0fffff", config)
    return "gAAAAAC" + req


# ─── Token Management ────────────────────────────────────────────────────────

# Cached access token. Re-populated on 401 via refresh_token (if configured).
_cached_access_token = None


class TokenExpiredError(RuntimeError):
    """ChatGPT rejected the access token (HTTP 401). Triggers auto-refresh."""


class RetryableError(RuntimeError):
    """Transient upstream failure (429, 5xx, timeout) worth retrying.

    status: HTTP status that caused it (0 = transport/network error).
    """

    def __init__(self, status: int, msg: str):
        super().__init__(msg)
        self.status = status


def refresh_token_to_access(refresh_token: str) -> str:
    url = "https://auth0.openai.com/oauth/token"
    data = {
        "redirect_uri": "com.openai.chat://auth0.openai.com/ios/com.openai.chat/callback",
        "grant_type": "refresh_token",
        "client_id": "pdlLIX2Y72MIl2rhLhTE9VV9bN905kB",
        "refresh_token": refresh_token,
    }
    try:
        session = cffi_requests.Session(impersonate=CONFIG.get("impersonate", "safari15_3"))
        r = session.post(url, data=data, timeout=15)
        r.raise_for_status()
        token = r.json()["access_token"]
        session.close()
        log("Access token refreshed OK")
        return token
    except Exception as e:
        log(f"Token refresh failed: {e}")
        raise


def get_access_token(force_refresh: bool = False):
    """Return cached access token; refresh from refresh_token when asked.

    Priority:
      1. Cached token (unless force_refresh)
      2. refresh_token in config (exchange → cache)
      3. access_token in config (use as-is)
    """
    global _cached_access_token
    if not force_refresh and _cached_access_token:
        return _cached_access_token

    at = CONFIG.get("access_token")
    rt = CONFIG.get("refresh_token")

    if force_refresh and rt:
        _cached_access_token = refresh_token_to_access(rt)
        return _cached_access_token
    if at:
        _cached_access_token = at
        return _cached_access_token
    if rt:
        _cached_access_token = refresh_token_to_access(rt)
        return _cached_access_token
    return None


def is_retryable_status(status: int) -> bool:
    """True if the upstream HTTP status is worth retrying."""
    return status == 429 or 500 <= status <= 599


# ─── ChatGPT Backend ─────────────────────────────────────────────────────────

def _make_session():
    """Create a curl_cffi session with browser impersonation."""
    imp = CONFIG.get("impersonate", "safari15_3")
    proxy = CONFIG.get("proxy")
    session = cffi_requests.Session(impersonate=imp, proxy=proxy)
    return session


def _base_headers(access_token):
    h = {
        "accept": "*/*",
        "accept-language": "en-US,en;q=0.9",
        "content-type": "application/json",
        "origin": "https://chatgpt.com",
        "referer": "https://chatgpt.com/",
        "user-agent": UA,
    }
    if access_token:
        h["authorization"] = f"Bearer {access_token}"
    return h


def init_page_cookies(session):
    """GET chatgpt.com/ to obtain essential cookies (oai-did, __cf_bm, etc)."""
    r = session.get(CONFIG["host_url"] + "/", headers={"user-agent": UA}, timeout=10)
    if r.status_code != 200:
        log(f"Page load warning: status {r.status_code}")
    # Also update DPL from page
    global _cached_scripts, _cached_dpl, _cached_dpl_time
    scripts = re.findall(r'<script[^>]+src="([^"]+)"', r.text)
    if scripts:
        _cached_scripts = scripts
    for src in scripts:
        m = re.search(r'/_next/static/([^/]+)/', src)
        if m:
            _cached_dpl = m.group(1)
            break
    if not _cached_dpl:
        m = re.search(r'data-build="([^"]+)"', r.text)
        if m:
            _cached_dpl = m.group(1)
    if _cached_dpl:
        _cached_dpl_time = int(time.time())
    return r.status_code


# ─── File Upload ────────────────────────────────────────────────────────────

MULTIMODAL_MIMES = {"image/jpeg", "image/webp", "image/png", "image/gif"}

MY_FILES_MIMES = {
    "text/x-php", "application/msword", "text/x-c", "text/html",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/json", "text/javascript", "application/pdf",
    "text/x-java", "text/x-tex", "text/x-typescript", "text/x-sh",
    "text/x-csharp", "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "text/x-c++", "application/x-latex", "text/markdown", "text/plain",
    "text/x-ruby", "text/x-script.python",
}

MIME_EXT_MAP = {
    "image/jpeg": ".jpg", "image/png": ".png", "image/gif": ".gif",
    "image/webp": ".webp", "application/pdf": ".pdf", "text/plain": ".txt",
    "text/markdown": ".md", "application/json": ".json", "text/html": ".html",
    "text/css": ".css", "text/xml": ".xml", "application/xml": ".xml",
    "application/msword": ".doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
    "text/x-python": ".py", "text/x-script.python": ".py", "text/x-sh": ".sh",
    "text/javascript": ".js", "text/x-java": ".java", "text/x-c": ".c",
    "text/x-c++": ".cpp", "text/x-csharp": ".cs", "text/x-ruby": ".rb",
    "text/x-tex": ".tex", "application/x-latex": ".latex",
    "application/zip": ".zip", "application/x-zip-compressed": ".zip",
    "application/x-tar": ".tar", "application/x-gzip": ".gz",
    "audio/mpeg": ".mp3", "audio/wav": ".wav", "video/mp4": ".mp4",
    "text/csv": ".csv", "application/rtf": ".rtf",
}


def _mime_from_name(filename):
    ext = os.path.splitext(filename or "")[1].lower()
    rev = {v: k for k, v in MIME_EXT_MAP.items() if v.startswith(ext)}
    if rev:
        return max(rev)  # longest matching extension (e.g. .docx over .doc)
    return "application/octet-stream"


def _get_file_extension(mime_type):
    return MIME_EXT_MAP.get(mime_type, "")


def _determine_use_case(mime_type):
    if mime_type in MULTIMODAL_MIMES:
        return "multimodal"
    if mime_type in MY_FILES_MIMES:
        return "my_files"
    return "ace_upload"


def _parse_data_url(data_url):
    """data:<mime>;base64,<payload> → (bytes, mime)."""
    m = re.match(r"data:([^;,]+)?(;base64)?,(.*)", data_url, re.S)
    if not m:
        return None, None
    mime, is_b64, payload = m.group(1), m.group(2), m.group(3)
    if is_b64:
        return pybase64.b64decode(payload), (mime or "application/octet-stream")
    return payload.encode("utf-8"), (mime or "text/plain")


def _fetch_file_content(source, session):
    """source = data: URL or http(s) URL → (bytes, mime_type)."""
    if isinstance(source, str) and source.startswith("data:"):
        content, mime = _parse_data_url(source)
        if content is None:
            raise RuntimeError("invalid data: URL in image_url")
        return content, mime
    r = session.get(source, timeout=60)
    if r.status_code != 200:
        raise RuntimeError(f"failed to fetch file URL {source[:80]}: HTTP {r.status_code}")
    mime = r.headers.get("Content-Type", "").split(";")[0].strip() or "application/octet-stream"
    return r.content, mime


def _get_image_size(file_content):
    if not HAS_PIL:
        raise RuntimeError("Pillow (pillow) required for image uploads — pip install pillow")
    with Image.open(io.BytesIO(file_content)) as img:
        return img.width, img.height


# ─── Local File Parsing (fallback when upload/read is unavailable) ──────────

TEXT_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".css", ".js",
    ".py", ".sh", ".log", ".ini", ".yaml", ".yml", ".tex", ".rtf",
}


def _decode_bytes(file_content, mime_type):
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_content.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return file_content.decode("utf-8", errors="replace")


def _extract_pdf_text(file_content):
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            return None, "PDF parsing needs pypdf: pip install pypdf"
    try:
        reader = PdfReader(io.BytesIO(file_content))
        pages = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                pages.append(t.strip())
    except Exception as e:
        return None, f"PDF parse error: {e}"
    if not pages:
        return None, "PDF has no extractable text (scanned/image PDF not supported)"
    return "\n\n".join(pages), None


def _extract_docx_text(file_content):
    try:
        from docx import Document
    except ImportError:
        return None, "DOCX parsing needs python-docx: pip install python-docx"
    try:
        doc = Document(io.BytesIO(file_content))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
    except Exception as e:
        return None, f"DOCX parse error: {e}"
    if not parts:
        return None, "DOCX has no extractable text"
    return "\n".join(parts), None


def _extract_xlsx_text(file_content):
    try:
        import openpyxl
    except ImportError:
        return None, "XLSX parsing needs openpyxl: pip install openpyxl"
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(x.strip() for x in cells):
                    rows.append("\t".join(cells))
            if rows:
                parts.append(f"[Sheet: {ws.title}]\n" + "\n".join(rows))
    except Exception as e:
        return None, f"XLSX parse error: {e}"
    if not parts:
        return None, "XLSX has no readable rows"
    return "\n\n".join(parts), None


def parse_file_text(file_content, mime_type, filename):
    """Extract text from a file locally. Returns (text, error).

    Images return (None, None) — they must go through the upload flow.
    Unknown binaries return (None, "unsupported...").
    """
    if not file_content:
        return None, "empty file"
    mime = (mime_type or "").lower()
    ext = os.path.splitext(filename or "")[1].lower()

    if mime.startswith("image/"):
        return None, None
    if mime.startswith("text/") or ext in TEXT_EXTENSIONS:
        return _decode_bytes(file_content, mime), None
    if mime == "application/pdf" or ext == ".pdf":
        return _extract_pdf_text(file_content)
    if mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",) or ext == ".docx":
        return _extract_docx_text(file_content)
    if mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",) or ext == ".xlsx":
        return _extract_xlsx_text(file_content)
    if mime == "application/msword" or ext == ".doc":
        return None, ".doc (old Word) not supported — convert to .docx or .pdf first"

    # Unknown type: inline if it looks like readable text
    try:
        decoded = _decode_bytes(file_content, mime)
        sample = decoded[:2000]
        printable = sum(1 for ch in sample if ch.isprintable() or ch in "\n\t\r")
        if sample and printable / len(sample) > 0.9:
            return decoded, None
    except Exception:
        pass
    return None, f"unsupported file type: {mime or ext or 'unknown'}"


def _create_upload_url(session, access_token, file_name, file_size, use_case, oai_device_id):
    url = CONFIG["host_url"] + "/backend-api/files"
    headers = _base_headers(access_token)
    headers["oai-device-id"] = oai_device_id
    body = {
        "file_name": file_name,
        "file_size": file_size,
        "reset_rate_limits": False,
        "timezone_offset_min": -480,
        "use_case": use_case,
    }
    r = session.post(url, headers=headers, json=body, timeout=15)
    if r.status_code == 401:
        raise TokenExpiredError("files: 401 unauthorized (token expired)")
    if r.status_code != 200:
        raise RuntimeError(f"files: HTTP {r.status_code} {r.text[:300]}")
    res = r.json()
    return res.get("file_id"), res.get("upload_url")


def _upload_blob(session, upload_url, file_content, mime_type):
    """Upload bytes to the presigned Azure Blob URL. No auth headers here."""
    headers = {
        "accept": "application/json, text/plain, */*",
        "content-type": mime_type,
        "x-ms-blob-type": "BlockBlob",
        "x-ms-version": "2020-04-08",
    }
    r = session.put(upload_url, headers=headers, data=file_content, timeout=60)
    if r.status_code != 201:
        raise RuntimeError(f"blob upload: HTTP {r.status_code} {r.text[:300]}")


def _register_upload(session, access_token, file_id):
    """POST /files/{id}/uploaded — registers the blob upload so ChatGPT can find it."""
    url = CONFIG["host_url"] + f"/backend-api/files/{file_id}/uploaded"
    headers = _base_headers(access_token)
    headers["oai-device-id"] = str(uuid.uuid4())
    r = session.post(url, headers=headers, json={}, timeout=10)
    if r.status_code == 401:
        raise TokenExpiredError("files/uploaded: 401 unauthorized (token expired)")
    if r.status_code != 200:
        raise RuntimeError(f"files/uploaded: HTTP {r.status_code} {r.text[:300]}")
    return r.json()


def _confirm_upload(session, access_token, file_id):
    """Poll /files/{id} until retrieval_index_status == success (non-image files)."""
    url = CONFIG["host_url"] + f"/backend-api/files/{file_id}"
    headers = _base_headers(access_token)
    headers["oai-device-id"] = str(uuid.uuid4())
    deadline = time.time() + 30
    while time.time() < deadline:
        try:
            r = session.get(url, headers=headers, timeout=10)
        except Exception as e:
            log(f"File status poll error: {e}")
            time.sleep(1)
            continue
        if r.status_code == 200:
            res = r.json()
            status = res.get("retrieval_index_status", "")
            if status == "success":
                log(f"File {file_id} indexed")
                return True
            if status == "failed":
                log(f"File {file_id} indexing FAILED")
                return False
        time.sleep(1)
    return False


def upload_file(session, access_token, file_content, mime_type, oai_device_id):
    """Full upload flow → file_meta dict (or None). Mirrors the ChatGPT web client."""
    if not file_content or not mime_type:
        return None

    width = height = None
    if mime_type.startswith("image/"):
        try:
            width, height = _get_image_size(file_content)
        except Exception as e:
            log(f"Image size unavailable, falling back to text/plain: {e}")
            mime_type = "text/plain"

    file_size = len(file_content)
    file_name = f"{uuid.uuid4()}{_get_file_extension(mime_type)}"
    use_case = _determine_use_case(mime_type)

    file_id, upload_url = _create_upload_url(session, access_token, file_name, file_size, use_case, oai_device_id)
    if not file_id or not upload_url:
        raise RuntimeError("no file_id/upload_url from /files")

    _upload_blob(session, upload_url, file_content, mime_type)

    # Register the upload so ChatGPT can find the file
    _register_upload(session, access_token, file_id)

    # Confirm indexing for retrievable files (skip ace_upload — raw uploads don't index)
    if use_case != "ace_upload":
        _confirm_upload(session, access_token, file_id)

    return {
        "file_id": file_id,
        "file_name": file_name,
        "size_bytes": file_size,
        "mime_type": mime_type,
        "width": width,
        "height": height,
        "use_case": use_case,
    }


def get_chat_requirements(session, access_token):
    """POST /backend-api/sentinel/chat-requirements."""
    url = CONFIG["host_url"] + "/backend-api/sentinel/chat-requirements"
    headers = _base_headers(access_token)
    headers["oai-device-id"] = str(uuid.uuid4())

    config = _get_pow_config(UA)
    p = _get_requirements_token(config)

    r = session.post(url, headers=headers, json={"p": p}, timeout=15)
    if r.status_code == 401:
        raise TokenExpiredError("chat-requirements: 401 unauthorized (token expired)")
    if is_retryable_status(r.status_code):
        raise RetryableError(r.status_code, f"chat-requirements: HTTP {r.status_code}")
    if r.status_code != 200:
        raise RuntimeError(f"chat-requirements failed: {r.status_code} {r.text[:300]}")
    return r.json(), config


def _upload_attachment(session, access_token, file_content, mime_type, oai_device_id, display_name=None):
    """Upload a file via the /files flow → attachments entry dict."""
    file_meta = upload_file(session, access_token, file_content, mime_type, oai_device_id)
    if not file_meta:
        raise RuntimeError("file upload failed")
    att = {
        "id": file_meta["file_id"],
        "size": file_meta["size_bytes"],
        "name": display_name or file_meta["file_name"],
        "mime_type": file_meta["mime_type"],
    }
    if file_meta.get("width") is not None:
        att["width"] = file_meta["width"]
        att["height"] = file_meta["height"]
    return att


def open_conversation(session, access_token, chat_token, proof_token, messages, model_slug, oai_device_id):
    """POST /backend-api/conversation (stream=True). Returns the response object.

    Raises TokenExpiredError on 401, RetryableError on 429/5xx, RuntimeError otherwise.
    The response body is NOT consumed here — caller iterates resp.iter_lines().
    Splitting open vs read lets us retry/refresh before any SSE bytes go to the client.
    """
    url = CONFIG["host_url"] + "/backend-api/conversation"
    headers = _base_headers(access_token)
    headers["accept"] = "text/event-stream"
    headers["oai-device-id"] = oai_device_id
    if chat_token:
        headers["openai-sentinel-chat-requirements-token"] = chat_token
    if proof_token:
        headers["openai-sentinel-proof-token"] = proof_token

    chat_messages = []
    file_mode = CONFIG.get("file_mode", "parse")
    for msg in messages:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        if isinstance(content, list):
            parts = []
            attachments = []
            for c in content:
                ctype = c.get("type")
                if ctype in ("text", "input_text"):
                    parts.append(c.get("text", ""))
                elif ctype == "image_url":
                    img_url = c.get("image_url", {}).get("url")
                    if not img_url:
                        continue
                    file_content, mime_type = _fetch_file_content(img_url, session)
                    if file_mode == "parse" and not (mime_type or "").startswith("image/"):
                        text, err = parse_file_text(file_content, mime_type, "")
                        if text is not None:
                            parts.append(f"[File attachment]\n{text.strip()}")
                            log("Inlined non-image attachment via image_url (parse mode)")
                            continue
                        if err:
                            log(f"Parse failed for image_url content: {err}")
                    # image (or upload mode): real multimodal upload
                    if (mime_type or "").startswith("image/"):
                        file_meta = upload_file(session, access_token, file_content, mime_type, oai_device_id)
                        if not file_meta:
                            raise RuntimeError("image upload failed")
                        fid = file_meta["file_id"]
                        parts.append({
                            "content_type": "image_asset_pointer",
                            "asset_pointer": f"file-service://{fid}",
                            "size_bytes": file_meta["size_bytes"],
                            "width": file_meta.get("width"),
                            "height": file_meta.get("height"),
                        })
                        att = {
                            "id": fid,
                            "size": file_meta["size_bytes"],
                            "name": file_meta["file_name"],
                            "mime_type": file_meta["mime_type"],
                        }
                        if file_meta.get("width") is not None:
                            att["width"] = file_meta["width"]
                            att["height"] = file_meta["height"]
                        attachments.append(att)
                    else:
                        # non-image in upload mode
                        attachments.append(_upload_attachment(
                            session, access_token, file_content, mime_type, oai_device_id))
                elif ctype in ("file", "input_file"):
                    # {type: file, file: {file_data|file_url, filename}}
                    finfo = c.get("file") or {}
                    source = finfo.get("file_data") or finfo.get("file_url") or finfo.get("url")
                    if not source:
                        continue
                    fname = finfo.get("filename") or ""
                    file_content, mime_type = _fetch_file_content(source, session)
                    if file_mode == "parse" and not (mime_type or "").startswith("image/"):
                        text, err = parse_file_text(file_content, mime_type, fname)
                        if text is not None:
                            label = fname or "attachment"
                            parts.append(f"[File: {label}]\n{text.strip()}")
                            log(f"Inlined file {label} ({len(file_content)} bytes) as text")
                            continue
                        if err:
                            log(f"Parse failed for {fname}: {err}")
                    # upload mode / image / unparseable → real upload
                    attachments.append(_upload_attachment(
                        session, access_token, file_content, mime_type, oai_device_id, fname))
            content_type = "multimodal_text" if len(attachments) else "text"
            if not parts:
                parts = [""]
            chat_messages.append({
                "id": str(uuid.uuid4()),
                "author": {"role": role},
                "content": {"content_type": content_type, "parts": parts},
                "metadata": {"attachments": attachments} if attachments else {},
            })
        else:
            chat_messages.append({
                "id": str(uuid.uuid4()),
                "author": {"role": role},
                "content": {"content_type": "text", "parts": [content]},
                "metadata": {},
            })

    body = {
        "action": "next",
        "client_contextual_info": {
            "is_dark_mode": False,
            "time_since_loaded": random.randint(50, 500),
            "page_height": random.randint(500, 1000),
            "page_width": random.randint(1000, 2000),
            "pixel_ratio": 1.5,
            "screen_height": random.randint(800, 1200),
            "screen_width": random.randint(1200, 2200),
        },
        "conversation_mode": {"kind": "primary_assistant"},
        "conversation_origin": None,
        "force_paragen": False,
        "force_paragen_model_slug": "",
        "force_rate_limit": False,
        "force_use_sse": True,
        "history_and_training_disabled": CONFIG.get("history_disabled", True),
        "messages": chat_messages,
        "model": model_slug,
        "paragen_cot_summary_display_override": "allow",
        "paragen_stream_type_override": None,
        "parent_message_id": str(uuid.uuid4()),
        "reset_rate_limits": False,
        "suggestions": [],
        "supported_encodings": [],
        "system_hints": [],
        "timezone": "America/Los_Angeles",
        "timezone_offset_min": -480,
        "variant_purpose": "comparison_implicit",
        "websocket_request_id": str(uuid.uuid4()),
    }

    resp = session.post(url, headers=headers, json=body, timeout=CONFIG["request_timeout_sec"], stream=True)
    if resp.status_code == 401:
        resp.close()
        raise TokenExpiredError("conversation: 401 unauthorized (token expired)")
    if is_retryable_status(resp.status_code):
        resp.close()
        raise RetryableError(resp.status_code, f"conversation: HTTP {resp.status_code}")
    if resp.status_code != 200:
        error_body = resp.text[:500] if hasattr(resp, 'text') else "unknown"
        resp.close()
        raise RuntimeError(f"conversation failed: {resp.status_code} {error_body}")
    return resp


def run_conversation(session, access_token, chat_token, proof_token, messages, model_slug, oai_device_id, retry_attempts=None, retry_delay=None):
    """Open the conversation with retry-on-transient errors.

    Returns a response whose iter_lines() yields the SSE stream. Caller consumes it.
    TokenExpiredError (401) is intentionally NOT handled here — it propagates so the
    caller can refresh the access token and redo the whole pipeline (fresh chat token too).
    """
    attempts = retry_attempts if retry_attempts is not None else CONFIG.get("retry_attempts", 3)
    delay = retry_delay if retry_delay is not None else CONFIG.get("retry_delay_sec", 2)

    last_err = None
    for attempt in range(attempts):
        try:
            return open_conversation(session, access_token, chat_token, proof_token, messages, model_slug, oai_device_id)
        except RetryableError as e:
            last_err = e
            if attempt < attempts - 1:
                log(f"Transient error (HTTP {e.status}), retry {attempt + 1}/{attempts}")
                time.sleep(delay)
                continue
            raise
        except Exception as e:
            if isinstance(e, TokenExpiredError):
                raise  # let the caller handle 401/refresh
            last_err = e
            # Retry transport-level errors (network drop, timeout) too.
            if attempt < attempts - 1:
                log(f"Transport error, retry {attempt + 1}/{attempts}: {e}")
                time.sleep(delay)
                continue
            raise

    raise last_err if last_err else RuntimeError("conversation failed")


# ─── SSE Parsing ─────────────────────────────────────────────────────────────

def _decode_line(line):
    if isinstance(line, bytes):
        return line.decode("utf-8", errors="replace")
    return line or ""


def parse_sse_stream(response_lines, model_name):
    """Parse ChatGPT SSE → OpenAI-format chunks. Yields SSE strings."""
    chat_id = f"chatcmpl-{''.join(random.choice(string.ascii_letters + string.digits) for _ in range(29))}"
    created = int(time.time())

    # First chunk: role
    yield f"data: {json.dumps({'id': chat_id, 'object': 'chat.completion.chunk', 'created': created, 'model': model_name, 'choices': [{'index': 0, 'delta': {'role': 'assistant', 'content': ''}, 'logprobs': None, 'finish_reason': None}]})}\n\n"

    prev_text = ""
    end = False

    for raw in response_lines:
        if end:
            break
        line = _decode_line(raw)
        if not line.startswith("data: "):
            continue
        data_str = line[6:]
        if data_str == "[DONE]":
            break
        try:
            data = json.loads(data_str)
        except json.JSONDecodeError:
            continue

        message = data.get("message", {})
        if not message:
            if data.get("error"):
                yield f"data: {json.dumps({'error': data['error']})}\n\n"
                break
            continue

        role = message.get("author", {}).get("role")
        if role in ("user", "system"):
            continue

        status = message.get("status")
        content = message.get("content", {})

        if status == "in_progress":
            ct = content.get("content_type")
            if ct == "text":
                parts = content.get("parts", [])
                if parts:
                    part = parts[0]
                    if isinstance(part, str) and len(part) > len(prev_text):
                        delta = part[len(prev_text):]
                        prev_text = part
                        yield f"data: {json.dumps({'id': chat_id, 'object': 'chat.completion.chunk', 'created': created, 'model': model_name, 'choices': [{'index': 0, 'delta': {'content': delta}, 'logprobs': None, 'finish_reason': None}]})}\n\n"

        elif status == "finished_successfully":
            if content.get("content_type") == "text":
                parts = content.get("parts", [])
                if parts and isinstance(parts[0], str):
                    final_part = parts[0]
                    delta = final_part[len(prev_text):] if len(final_part) > len(prev_text) else ""
                    if delta:
                        yield f"data: {json.dumps({'id': chat_id, 'object': 'chat.completion.chunk', 'created': created, 'model': model_name, 'choices': [{'index': 0, 'delta': {'content': delta}, 'logprobs': None, 'finish_reason': None}]})}\n\n"
            # Final chunk
            yield f"data: {json.dumps({'id': chat_id, 'object': 'chat.completion.chunk', 'created': created, 'model': model_name, 'choices': [{'index': 0, 'delta': {}, 'logprobs': None, 'finish_reason': 'stop'}]})}\n\n"
            end = True

    if not end:
        yield f"data: {json.dumps({'id': chat_id, 'object': 'chat.completion.chunk', 'created': created, 'model': model_name, 'choices': [{'index': 0, 'delta': {}, 'logprobs': None, 'finish_reason': 'stop'}]})}\n\n"

    yield "data: [DONE]\n\n"


def collect_full_text(response_lines):
    """Parse SSE stream → (full_text, model_slug)."""
    full_text = ""
    for raw in response_lines:
        line = _decode_line(raw)
        if not line.startswith("data: "):
            continue
        data_str = line[6:]
        if data_str == "[DONE]":
            break
        try:
            data = json.loads(data_str)
        except json.JSONDecodeError:
            continue
        message = data.get("message", {})
        if not message:
            continue
        role = message.get("author", {}).get("role")
        if role in ("user", "system"):
            continue
        status = message.get("status")
        content = message.get("content", {})
        if status == "finished_successfully" and content.get("content_type") == "text":
            parts = content.get("parts", [])
            if parts and isinstance(parts[0], str):
                full_text = parts[0]
                break
        elif status == "in_progress" and content.get("content_type") == "text":
            parts = content.get("parts", [])
            if parts and isinstance(parts[0], str):
                full_text = parts[0]
    return full_text


# ─── HTTP Handler ────────────────────────────────────────────────────────────

class ChatGPTHandler(BaseHTTPRequestHandler):
    def log_message(self, fmt, *args):
        client_ip = self.client_address[0] if self.client_address else "-"
        log(f"{client_ip} {fmt % args}")

    def send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode()
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _authorized(self):
        keys = CONFIG.get("api_keys") or []
        if not keys:
            return True
        auth = self.headers.get("Authorization", "")
        if auth.startswith("Bearer ") and auth[7:] in keys:
            return True
        if self.headers.get("x-api-key", "") in keys:
            return True
        return False

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "*")
        self.end_headers()

    def do_GET(self):
        try:
            if self.path.startswith("/v1") and not self._authorized():
                self.send_json({"error": {"message": "invalid api key"}}, 401)
                return
            if self.path == "/v1/models":
                self.send_json({"object": "list", "data": [
                    {"id": n, "object": "model", "created": 1700000000,
                     "owned_by": "openai", "description": c["desc"]}
                    for n, c in MODELS.items()
                ]})
            elif self.path == "/" or self.path == "/health":
                self.send_json({"status": "ok", "version": __version__,
                                "models": list(MODELS.keys()),
                                "token_configured": bool(CONFIG.get("access_token") or CONFIG.get("refresh_token"))})
            else:
                self.send_json({"error": "not found"}, 404)
        except (BrokenPipeError, ConnectionResetError):
            pass
        except Exception as e:
            log(f"GET error: {e}")

    def do_POST(self):
        try:
            if self.path.startswith("/v1") and not self._authorized():
                self.send_json({"error": {"message": "invalid api key"}}, 401)
                return
            length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(length) if length else b""
            if self.path == "/v1/chat/completions":
                self.handle_chat(body)
            else:
                self.send_json({"error": "not found"}, 404)
        except (BrokenPipeError, ConnectionResetError):
            pass
        except Exception as e:
            log(f"POST error: {e}")
            try:
                self.send_json({"error": {"message": str(e)}}, 500)
            except:
                pass

    def handle_chat(self, body: bytes):
        req = json.loads(body)
        model_name = req.get("model", CONFIG["default_model"])
        model_cfg = MODELS.get(model_name)
        if not model_cfg:
            for k, v in MODELS.items():
                if k.startswith(model_name) or model_name.startswith(k):
                    model_name, model_cfg = k, v
                    break
        if not model_cfg:
            self.send_json({"error": {"message": f"Unknown model: {model_name}"}}, 400)
            return
        model_slug = model_cfg["slug"]

        messages = req.get("messages", [])
        if not messages:
            self.send_json({"error": {"message": "empty messages"}}, 400)
            return

        stream = req.get("stream", False)
        log(f"Chat: model={model_name} slug={model_slug} stream={stream} msgs={len(messages)}")

        # Get access token
        try:
            access_token = get_access_token()
        except Exception as e:
            self.send_json({"error": {"message": f"auth error: {e}"}}, 401)
            return
        if not access_token:
            self.send_json({"error": {"message": "No access_token or refresh_token configured"}}, 401)
            return

        # Create session with browser impersonation
        session = _make_session()

        try:
            # Retry loop covers the whole pipeline: 401 → auto-refresh token & redo;
            # transient errors (429/5xx) → backoff & redo.
            attempts = CONFIG.get("retry_attempts", 3)
            delay = CONFIG.get("retry_delay_sec", 2)
            last_err = None
            resp = None
            oai_device_id = str(uuid.uuid4())
            for attempt in range(attempts):
                try:
                    # Step 1: Load page for cookies
                    init_page_cookies(session)

                    # Step 2: Get chat requirements
                    req_data, pow_config = get_chat_requirements(session, access_token)
                    chat_token = req_data.get("token")
                    if not chat_token:
                        raise RuntimeError(f"No chat token: {str(req_data)[:300]}")

                    # Step 3: Solve PoW
                    proof_token = None
                    pow_data = req_data.get("proofofwork", {})
                    if pow_data.get("required"):
                        pow_seed = pow_data.get("seed", "")
                        pow_diff = pow_data.get("difficulty", CONFIG["pow_difficulty"])
                        log(f"PoW: seed={pow_seed[:20]}... diff={pow_diff}")
                        proof_token, solved = _get_answer_token(pow_seed, pow_diff, pow_config)
                        log(f"PoW solved: {solved}")
                    else:
                        log("PoW not required")

                    # Step 4: Send conversation (retry + 401-refresh inside)
                    resp = run_conversation(
                        session, access_token, chat_token, proof_token,
                        messages, model_slug, oai_device_id,
                        retry_attempts=1, retry_delay=delay,
                    )
                    break
                except TokenExpiredError as e:
                    # 401 anywhere in the pipeline → refresh token once, redo everything.
                    if CONFIG.get("refresh_token"):
                        log("401 → refreshing access token via refresh_token")
                        try:
                            access_token = get_access_token(force_refresh=True)
                        except Exception as re:
                            raise RuntimeError(f"token refresh failed: {re}") from e
                        log("Token refreshed, restarting pipeline")
                        continue  # redo from step 1 with new token
                    self.send_json({
                        "error": {"message": "access token expired and no refresh_token configured — update config.json"},
                        "hint": "Get a fresh accessToken from https://chatgpt.com/api/auth/session, or set refresh_token",
                    }, 401)
                    return
                except RetryableError as e:
                    last_err = e
                    if attempt < attempts - 1:
                        log(f"Transient error (HTTP {e.status}), retry {attempt + 1}/{attempts}")
                        time.sleep(delay)
                        continue
                    raise
                except Exception as e:
                    last_err = e
                    if attempt < attempts - 1:
                        log(f"Pipeline error, retry {attempt + 1}/{attempts}: {e}")
                        time.sleep(delay)
                        continue
                    raise

            if resp is None:
                raise last_err if last_err else RuntimeError("pipeline failed")

            sse_lines = resp.iter_lines()

            if stream:
                self.send_response(200)
                self.send_header("Content-Type", "text/event-stream")
                self.send_header("Cache-Control", "no-cache")
                self.send_header("Access-Control-Allow-Origin", "*")
                self.end_headers()
                for chunk in parse_sse_stream(sse_lines, model_name):
                    self.wfile.write(chunk.encode())
                    self.wfile.flush()
            else:
                full_text = collect_full_text(sse_lines)
                chat_id = f"chatcmpl-{''.join(random.choice(string.ascii_letters + string.digits) for _ in range(29))}"
                if not full_text.strip():
                    self.send_json({"error": {"message": "Empty response from ChatGPT"}}, 502)
                    return
                self.send_json({
                    "id": chat_id,
                    "object": "chat.completion",
                    "created": int(time.time()),
                    "model": model_name,
                    "choices": [{
                        "index": 0,
                        "message": {"role": "assistant", "content": full_text},
                        "logprobs": None,
                        "finish_reason": "stop",
                    }],
                    "usage": {"prompt_tokens": 0, "completion_tokens": len(full_text) // 4, "total_tokens": len(full_text) // 4},
                })

        except Exception as e:
            log(f"Conversation error: {e}")
            try:
                self.send_json({"error": {"message": f"upstream error: {e}"}}, 502)
            except:
                pass
        finally:
            session.close()


# ─── Main ────────────────────────────────────────────────────────────────────

def load_config(path):
    if path and os.path.exists(path):
        with open(path) as f:
            CONFIG.update(json.load(f))
        log(f"Config loaded: {path}")


def main():
    parser = argparse.ArgumentParser(description="ChatGPT Web to OpenAI API")
    parser.add_argument("--port", type=int, default=None)
    parser.add_argument("--config", type=str, default=None)
    parser.add_argument("--access-token", type=str, default=None)
    parser.add_argument("--refresh-token", type=str, default=None)
    parser.add_argument("--proxy", type=str, default=None)
    parser.add_argument("--version", action="version", version=f"chatgpt-web2api {__version__}")
    args = parser.parse_args()

    config_path = args.config or os.environ.get("CHATGPT_WEB2API_CONFIG")
    if not config_path:
        for p in ["./config.json", os.path.expanduser("~/.config/chatgpt-web2api/config.json")]:
            if os.path.exists(p):
                config_path = p
                break
    load_config(config_path)

    if args.port:
        CONFIG["port"] = args.port
    if args.access_token:
        CONFIG["access_token"] = args.access_token
    if args.refresh_token:
        CONFIG["refresh_token"] = args.refresh_token
    if args.proxy:
        CONFIG["proxy"] = args.proxy

    if not HAS_CFFI:
        print("Error: curl_cffi is required. Install with: pip install curl_cffi")
        sys.exit(1)

    class ThreadedServer(ThreadingMixIn, HTTPServer):
        daemon_threads = True
        allow_reuse_address = True

    port = CONFIG["port"]
    server = ThreadedServer((CONFIG["host"], port), ChatGPTHandler)
    print(f"chatgpt-web2api v{__version__}")
    print(f"  Listening: http://0.0.0.0:{port}")
    print(f"  Base URL:  http://localhost:{port}/v1")
    print(f"  Models:    {', '.join(MODELS.keys())}")
    print(f"  Token:     {'yes' if CONFIG.get('access_token') or CONFIG.get('refresh_token') else 'none'}")
    print(f"  Proxy:     {CONFIG.get('proxy') or 'none'}")
    print(f"  Impersonate: {CONFIG.get('impersonate', 'safari15_3')}")
    print()
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopped.")
        server.shutdown()


if __name__ == "__main__":
    main()